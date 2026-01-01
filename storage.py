from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import pytz
from config import OUTPUT_FILE

def save_doc(records):
    """
    Save articles to a clean, well-formatted Word document.
    Uses bullet point summaries for easy scanning.
    All times displayed in AWST (Australian Western Standard Time).
    """
    document = Document()
    
    # Get current time in AWST
    awst = pytz.timezone('Australia/Perth')
    current_time_awst = datetime.now(awst)
    
    # Title
    title = document.add_heading("Digital Health News Summary - Australia", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Metadata
    meta = document.add_paragraph()
    meta_run = meta.add_run(
        f"Generated: {current_time_awst.strftime('%d %B %Y at %I:%M %p AWST')} | "
        f"Total Articles: {len(records)}"
    )
    meta_run.italic = True
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    document.add_paragraph()  # Single spacing
    
    # Add each article
    for i, record in enumerate(records, 1):
        # Article number and title
        heading = document.add_heading(f"Article {i}", level=2)
        
        # Title
        title_para = document.add_paragraph()
        title_para.add_run("Title: ").bold = True
        title_para.add_run(record.get('Title', 'Untitled'))
        
        # Summary
        summary_para = document.add_paragraph()
        summary_para.add_run("Summary:").bold = True
        
        # Add the bullet points (they should already be formatted from the AI)
        summary_text = record.get('Summary', 'Summary not available.')
        summary_content = document.add_paragraph(summary_text)
        
        # Source
        source_para = document.add_paragraph()
        source_para.add_run("Source: ").bold = True
        source_para.add_run(record.get('Link', 'N/A'))
        
        # Date - convert to AWST if possible
        date_para = document.add_paragraph()
        date_para.add_run("Published: ").bold = True
        
        original_date = record.get('Date', 'N/A')
        awst_date = convert_to_awst(original_date)
        date_para.add_run(awst_date)
        
        # Separator between articles (only if not the last article)
        if i < len(records):
            document.add_paragraph("─" * 80)
    
    # Save with timestamp
    timestamp = current_time_awst.strftime("%Y%m%d_%H%M%S")
    file_name = OUTPUT_FILE.replace(".docx", f"_{timestamp}.docx")
    
    document.save(file_name)
    print(f"\n✅ Saved {len(records)} articles to {file_name}")
    print(f"📄 Format: Clean bullet point summaries")
    print(f"🕐 All times shown in AWST (Australian Western Standard Time)")

def convert_to_awst(date_string: str) -> str:
    """
    Convert date string to AWST timezone.
    Handles multiple common date formats.
    """
    if not date_string or date_string == 'N/A':
        return 'N/A'
    
    awst = pytz.timezone('Australia/Perth')
    
    # Common date formats from news feeds
    date_formats = [
        "%Y-%m-%dT%H:%M:%S%z",           # ISO format with timezone
        "%Y-%m-%dT%H:%M:%SZ",             # ISO format UTC
        "%a, %d %b %Y %H:%M:%S %Z",       # RSS format
        "%a, %d %b %Y %H:%M:%S %z",       # RSS format with timezone
        "%Y-%m-%d %H:%M:%S",              # Simple datetime
        "%Y-%m-%d",                       # Date only
    ]
    
    for fmt in date_formats:
        try:
            # Parse the date
            if '%z' in fmt or '%Z' in fmt:
                # Has timezone info
                dt = datetime.strptime(date_string.strip(), fmt)
                if dt.tzinfo is None:
                    # If parsing removed timezone, assume UTC
                    dt = pytz.UTC.localize(dt)
            else:
                # No timezone info, assume UTC
                dt = datetime.strptime(date_string.strip(), fmt)
                dt = pytz.UTC.localize(dt)
            
            # Convert to AWST
            dt_awst = dt.astimezone(awst)
            return dt_awst.strftime('%d %B %Y at %I:%M %p AWST')
            
        except ValueError:
            continue
    
    # If no format worked, return original with AWST note
    return f"{date_string} (timezone unknown)"